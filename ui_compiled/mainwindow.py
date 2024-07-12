# This Python file uses the following encoding: utf-8
import sys
from PySide6.QtWidgets import QApplication, QMainWindow

from main_menu_ui import Ui_Dialog
from docx import Document
def debug_replace_no_table(template_path, output_path,data):# no table
    doc=Document(template_path)
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
               for run in paragraph.runs:
                   run.text=run.text.replace(key, value)
    doc.save(output_path)
def debug_replace(template_path, output_path,data):
    doc=Document(template_path)
    for table in doc.tables:
            table.cell(0,1).text = "datedata"
            table.cell(0, 1).text = "customerdata"
            table.cell(0, 1).text = "addres"
            table.cell(0, 1).text = "telef"
            table.cell(0, 1).text = "rabota"
            table.cell(0, 1).text = "imya"
            table.cell(0, 1).text = "various"
            table.cell(0, 1).text = "buydate"
            table.cell(0, 1).text = "problem"

    doc.save(output_path)
class MainWindow(QMainWindow):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.ui = Ui_Dialog()
        self.ui.setupUi(self)


if __name__ == "__main__":
    app = QApplication(sys.argv)
    widget = MainWindow()



    widget.show()
    sys.exit(app.exec())

